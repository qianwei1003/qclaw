from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# 标题
title = doc.add_heading('文档标题', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 元信息
meta = doc.add_paragraph()
meta.add_run('来源: ').bold = True
meta.add_run('[自动填充]')
meta.add_run('\n抓取时间: ').bold = True
meta.add_run('[自动填充]')

# 分隔线
doc.add_paragraph('---')

# 正文占位
doc.add_paragraph('[正文内容将插入此处]')

# 保存
doc.save('C:\\Users\\60597\\.qclaw\\workspace\\skills\\web-crawler-doc\\assets\\report-template.docx')
print('模板已创建')
