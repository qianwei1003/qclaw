#!/usr/bin/env python3
"""
网页爬虫脚本 - 单页爬取
支持 Markdown 和 Word 输出
"""

import argparse
import os
import re
import sys
from urllib.parse import urljoin, urlparse
from datetime import datetime

try:
    import requests
    from bs4 import BeautifulSoup
    from markdownify import markdownify as md
except ImportError:
    print("请先安装依赖: pip install requests beautifulsoup4 markdownify")
    sys.exit(1)


def fetch_page(url, headers=None, timeout=30):
    """获取网页内容"""
    default_headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    }
    if headers:
        default_headers.update(headers)
    
    try:
        response = requests.get(url, headers=default_headers, timeout=timeout)
        response.raise_for_status()
        response.encoding = response.apparent_encoding
        return response.text
    except Exception as e:
        print(f"获取页面失败: {e}")
        return None


def extract_content(html, selector=None):
    """提取正文内容"""
    soup = BeautifulSoup(html, 'html.parser')
    
    # 移除脚本和样式
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    
    # 如果指定了选择器
    if selector:
        content = soup.select_one(selector)
        if content:
            return content, soup.title.string if soup.title else "无标题"
    
    # 自动检测正文区域
    # 优先常见的正文容器
    content_selectors = [
        'article',
        '.article-content',
        '.content',
        '.post-content',
        '.entry-content',
        '#content',
        '.main-content',
        '[role="main"]'
    ]
    
    for sel in content_selectors:
        content = soup.select_one(sel)
        if content:
            return content, soup.title.string if soup.title else "无标题"
    
    #  fallback: 找最长文本的 div
    paragraphs = soup.find_all('p')
    if paragraphs:
        # 找包含最多段落的父元素
        parent_counts = {}
        for p in paragraphs:
            parent = p.find_parent(['div', 'section', 'article'])
            if parent:
                parent_counts[parent] = parent_counts.get(parent, 0) + 1
        
        if parent_counts:
            best_parent = max(parent_counts, key=parent_counts.get)
            return best_parent, soup.title.string if soup.title else "无标题"
    
    # 最后 fallback: 返回 body
    return soup.body or soup, soup.title.string if soup.title else "无标题"


def clean_title(title):
    """清理标题，用于文件名"""
    if not title:
        return "untitled"
    # 移除非法字符
    title = re.sub(r'[\\/*?:"<>|]', '', title)
    # 限制长度
    return title[:100].strip()


def save_markdown(content_html, title, url, save_path):
    """保存为 Markdown"""
    # 转换为 markdown
    markdown_content = md(str(content_html), heading_style="ATX")
    
    # 添加元信息
    header = f"""# {title}

> 来源: {url}
> 抓取时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

"""
    
    full_content = header + markdown_content
    
    # 确保目录存在
    os.makedirs(save_path, exist_ok=True)
    
    # 生成文件名
    filename = f"{clean_title(title)}_{datetime.now().strftime('%Y%m%d')}.md"
    filepath = os.path.join(save_path, filename)
    
    # 保存
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(full_content)
    
    print(f"✓ Markdown 已保存: {filepath}")
    return filepath


def save_word(content_html, title, url, save_path, template=None):
    """保存为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        return None
    
    # 创建文档
    if template and os.path.exists(template):
        doc = Document(template)
    else:
        doc = Document()
    
    # 添加标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加元信息
    meta = doc.add_paragraph()
    meta.add_run(f"来源: ").bold = True
    meta.add_run(url)
    meta.add_run(f"\n抓取时间: ").bold = True
    meta.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    
    doc.add_paragraph("---")
    
    # 解析内容并添加到文档
    soup = BeautifulSoup(str(content_html), 'html.parser')
    
    for elem in soup.descendants:
        if elem.name == 'h1':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'p':
            text = elem.get_text().strip()
            if text:
                doc.add_paragraph(text)
        elif elem.name == 'li':
            text = elem.get_text().strip()
            if text:
                doc.add_paragraph(text, style='List Bullet')
        elif elem.name == 'table':
            # 简单表格处理
            rows = elem.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell.get_text().strip()
    
    # 确保目录存在
    os.makedirs(save_path, exist_ok=True)
    
    # 生成文件名
    filename = f"{clean_title(title)}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(save_path, filename)
    
    # 保存
    doc.save(filepath)
    print(f"✓ Word 已保存: {filepath}")
    return filepath


def main():
    parser = argparse.ArgumentParser(description='网页爬虫 - 生成文档')
    parser.add_argument('--url', required=True, help='目标网址')
    parser.add_argument('--output', choices=['markdown', 'word', 'both'], default='markdown',
                       help='输出格式 (默认: markdown)')
    parser.add_argument('--selector', help='CSS 选择器，指定提取区域')
    parser.add_argument('--title', help='自定义文档标题')
    parser.add_argument('--save-to', default='./docs/', help='保存路径')
    parser.add_argument('--template', help='Word 模板文件路径')
    
    args = parser.parse_args()
    
    print(f"🌐 正在爬取: {args.url}")
    
    # 获取页面
    html = fetch_page(args.url)
    if not html:
        print("❌ 获取页面失败")
        return 1
    
    # 提取内容
    content, page_title = extract_content(html, args.selector)
    title = args.title or page_title
    
    print(f"📄 标题: {title}")
    
    # 生成文档
    saved_files = []
    
    if args.output in ['markdown', 'both']:
        md_path = save_markdown(content, title, args.url, args.save_to)
        if md_path:
            saved_files.append(md_path)
    
    if args.output in ['word', 'both']:
        docx_path = save_word(content, title, args.url, args.save_to, args.template)
        if docx_path:
            saved_files.append(docx_path)
    
    if saved_files:
        print(f"\n✅ 完成！共生成 {len(saved_files)} 个文件")
        for f in saved_files:
            print(f"   - {f}")
        return 0
    else:
        print("\n❌ 未生成任何文件")
        return 1


if __name__ == '__main__':
    sys.exit(main())
